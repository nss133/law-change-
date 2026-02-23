"""
DOCX 문서 생성 모듈
- 템플릿(소득세법 시행령 시행 안내) 형식 준수:
  - A4, 여백 12.7mm
  - 제목/메타: 14pt KoPub돋움체 Bold, 줄간격 1.65
  - 본문: 11pt KoPub돋움체 Light, 줄간격 2.2
"""

from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Tuple, Union

# 문서 표준 설정 (템플릿 기준)
LINE_SPACING_BODY = 2.2           # 본문 줄간격 (약 220%)
LINE_SPACING_TITLE = 1.65         # 제목/메타 줄간격
PARAGRAPH_SPACING_BEFORE = Pt(0)  # 문단 앞 간격
PARAGRAPH_SPACING_AFTER = Pt(0)   # 문단 뒤 간격


def _apply_body_format(paragraph):
    """본문 문단에 표준 형식 적용"""
    paragraph.paragraph_format.line_spacing = LINE_SPACING_BODY
    paragraph.paragraph_format.space_before = PARAGRAPH_SPACING_BEFORE
    paragraph.paragraph_format.space_after = PARAGRAPH_SPACING_AFTER


class DocxGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()
        self._toc_items = []

    def _setup_page(self):
        """용지 크기 및 여백 설정 (A4, 12.7mm)"""
        section = self.doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        margin = Mm(12.7)
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'KoPub돋움체_Pro Light'
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'KoPub돋움체_Pro Light')
        style.paragraph_format.line_spacing = LINE_SPACING_BODY

    def add_title(self, title_text: str):
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'KoPub돋움체_Pro Bold'
        title.paragraph_format.line_spacing = LINE_SPACING_TITLE
        title.paragraph_format.space_after = Pt(0)

    def add_metadata(self, enforcement_date: str, law_number: str, amendment_date: str, date: str = "25. 01.", dept: str = "법 무 팀"):
        # 시행일/법률번호 정보가 있을 때만 메타 정보 출력 (입법예고문 등은 생략)
        if enforcement_date or law_number or amendment_date:
            meta1 = self.doc.add_paragraph()
            meta1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = meta1.add_run(f"[시행 {enforcement_date}] [법률 제{law_number}호, {amendment_date}]")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'KoPub돋움체_Pro Bold'
            meta1.paragraph_format.line_spacing = LINE_SPACING_TITLE
            meta1.paragraph_format.space_before = PARAGRAPH_SPACING_BEFORE
            meta1.paragraph_format.space_after = PARAGRAPH_SPACING_AFTER

        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.add_run(date)
        _apply_body_format(date_p)

        dept_p = self.doc.add_paragraph()
        dept_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dept_p.add_run(dept)
        _apply_body_format(dept_p)

        self.doc.add_paragraph()

    def add_toc(self, items: List[Tuple[str, str]]):
        """목차 추가 (1. 개정 이유, 2. 주요내용 ...)"""
        toc_title = self.doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_title.add_run("목 차")
        run.bold = True
        run.font.size = Pt(12)
        toc_title.paragraph_format.space_before = Pt(12)
        toc_title.paragraph_format.space_after = Pt(12)

        for num, title in items:
            toc_p = self.doc.add_paragraph()
            toc_p.add_run(f"{num}. {title}")
            toc_p.paragraph_format.left_indent = Inches(0.3)
            _apply_body_format(toc_p)

        self.doc.add_paragraph()
        toc_p = self.doc.add_paragraph()  # 구분선 대신 빈 문단
        toc_p.paragraph_format.space_after = Pt(18)

    def add_section(self, number: str, title: str, content: Union[str, List[str]] = None, is_bold: bool = False):
        """섹션 추가. content가 리스트면 문단별로 분리 출력"""
        section = self.doc.add_paragraph()
        run = section.add_run(f"{number}. {title}")
        run.bold = True
        run.font.size = Pt(11)
        _apply_body_format(section)

        if content:
            if isinstance(content, list):
                for para_text in content:
                    if not para_text.strip():
                        continue
                    content_p = self.doc.add_paragraph()
                    content_run = content_p.add_run(para_text.strip())
                    if is_bold:
                        content_run.bold = True
                    content_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    content_p.paragraph_format.first_line_indent = Inches(0)
                    _apply_body_format(content_p)
            else:
                content_p = self.doc.add_paragraph()
                content_run = content_p.add_run(content.strip())
                if is_bold:
                    content_run.bold = True
                content_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _apply_body_format(content_p)

        self.doc.add_paragraph()

    def add_main_contents(self, contents: List[Dict] = None, intro_paragraphs: List[str] = None):
        """주요내용 - TXT 자료(intro_paragraphs)만 출력, 조문 내용 미포함"""
        if intro_paragraphs:
            for para_text in intro_paragraphs:
                if not para_text.strip():
                    continue
                p = self.doc.add_paragraph()
                p.add_run(para_text.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _apply_body_format(p)
            self.doc.add_paragraph()

        # TXT 자료에만 의존, 조문(법문) 내용은 출력하지 않음
        self.doc.add_paragraph()

    def add_impact_analysis(self, impact_text: str):
        impact_p = self.doc.add_paragraph()
        run = impact_p.add_run(impact_text)
        run.bold = True
        run.font.size = Pt(10)
        self.doc.add_paragraph()

    def add_comparison_table(self, comparison_data: List[Tuple]):
        if not comparison_data:
            return

        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = f"{comparison_data[0][0].split('[')[0].strip()}\n[개정 전]"
        header_cells[1].text = f"{comparison_data[0][1].split('[')[0].strip()}\n[개정 후]"

        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_background(cell, "D9D9D9")

        for old_text, new_text in comparison_data:
            row_cells = table.add_row().cells
            row_cells[0].text = old_text
            row_cells[1].text = new_text

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = LINE_SPACING_BODY
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        self.doc.add_paragraph()

    def _set_cell_background(self, cell, color: str):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)

    def save(self, filename: str):
        self.doc.save(filename)
        print(f"✅ 문서 저장 완료: {filename}")
